import re
import uuid
from copy import deepcopy

from .novel_utils import (
    _canonicalize_character_record,
    _character_activity_stats,
    _character_merge_key,
    _clip_context_text,
    _compact_text,
    _dedupe_text_lines,
    _infer_foreshadow_status,
    _lore_activity_stats,
    _merge_text_lines_without_duplicates,
    _normalize_name_list,
    _prepare_character_candidate,
    _record_alias_keys,
)

try:
    from docx import Document
    from docx.enum.text import WD_ALIGN_PARAGRAPH
except Exception:
    Document = None
    WD_ALIGN_PARAGRAPH = None


def _read_docx_text(path):
    if Document is None:
        raise RuntimeError("当前环境缺少 python-docx，无法导入 Word 文档。")
    doc = Document(path)
    lines = []
    for para in doc.paragraphs:
        text = str(para.text or "").strip()
        if text:
            lines.append(text)
    for table in doc.tables:
        for row in table.rows:
            cells = []
            for cell in row.cells:
                cell_text = "\n".join(
                    str(para.text or "").strip()
                    for para in cell.paragraphs
                    if str(para.text or "").strip()
                )
                if cell_text:
                    cells.append(cell_text)
            if cells:
                lines.append("\n".join(cells))
    return "\n".join(lines)


def _read_txt_text(path):
    last_error = None
    for encoding in ("utf-8", "utf-8-sig", "gb18030", "utf-16", "utf-16-le", "utf-16-be"):
        try:
            with open(path, "r", encoding=encoding) as f:
                return f.read()
        except UnicodeDecodeError as e:
            last_error = e
            continue
    if last_error:
        raise last_error
    with open(path, "r", encoding="utf-8") as f:
        return f.read()


def _write_docx_text(path, title, chapters, center_chapter_headings=False):
    if Document is None:
        raise RuntimeError("当前环境缺少 python-docx，无法导出 Word 文档。")
    doc = Document()
    title_para = doc.add_heading(str(title or "未命名小说"), 0)
    if center_chapter_headings and WD_ALIGN_PARAGRAPH is not None:
        title_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
    for index, chap in enumerate(chapters if isinstance(chapters, list) else [], 1):
        if not isinstance(chap, dict):
            continue
        text = str(chap.get("text", "") or "").strip()
        if not text:
            continue
        heading = str(chap.get("title", "") or f"第 {index} 章").strip()
        heading_para = doc.add_heading(heading, level=1)
        if center_chapter_headings and WD_ALIGN_PARAGRAPH is not None:
            heading_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
        for para in text.splitlines():
            para = para.strip()
            if para:
                doc.add_paragraph(para)
    doc.save(path)


def _unique_keep_order(values, limit=80):
    seen = set()
    out = []
    for value in values:
        value = str(value or "").strip()
        if not value or value in seen:
            continue
        seen.add(value)
        out.append(value)
        if len(out) >= limit:
            break
    return out


def _extract_import_candidates(text):
    text = str(text or "")
    compact = re.sub(r"\s+", "", text)
    common_words = {
        "他们", "我们", "你们", "自己", "这个", "那个", "什么", "没有", "不是", "已经", "突然", "只是",
        "时候", "地方", "事情", "眼前", "身后", "面前", "声音", "所有", "一个", "一种",
    }
    names = []
    for match in re.findall(r"[\u4e00-\u9fff]{2,4}", compact):
        if match in common_words:
            continue
        if match.endswith(("说道", "看着", "知道", "觉得", "起来", "下去", "里面", "这里", "那里")):
            continue
        names.append(match)

    place_keywords = "府|宫|城|镇|村|山|宗|门|阁|楼|院|司|部|军|营|帮|派|国|朝|殿|寺|谷|湖|江|河"
    lore = []
    for match in re.findall(rf"[\u4e00-\u9fff]{{2,8}}(?:{place_keywords})", compact):
        lore.append(match)
    for match in re.findall(r"[\u4e00-\u9fff]{1,6}(?:令|符|剑|刀|玉|珠|书|卷|印|冠|甲|丹|药)", compact):
        lore.append(match)

    foreshadows = []
    for sentence in re.split(r"[。！？!?；;]\s*", text):
        s = sentence.strip()
        if not s:
            continue
        if any(k in s for k in ("秘密", "真相", "失踪", "旧案", "线索", "疑点", "隐藏", "身份", "遗物", "谜")):
            foreshadows.append(s[:60])

    character_candidates = []
    lore_candidates = [{"name": name, "type": "其他", "description": "文档导入候选"} for name in _unique_keep_order(lore, 40)]
    for name in _unique_keep_order(names, 40):
        char_item, lore_item = _prepare_character_candidate(
            {"name": name, "role": "", "goal": "", "secret": "", "voice": "", "notes": "文档导入候选"},
            default_notes="文档导入候选",
        )
        if char_item:
            character_candidates.append(char_item)
        elif lore_item:
            lore_candidates.append(lore_item)

    return {
        "characters": character_candidates,
        "lore": lore_candidates,
        "foreshadows": [
            {
                "name": item,
                "status": _infer_foreshadow_status({"name": item, "description": item}, preserve_manual=False),
                "setup_chapter": "",
                "payoff_chapter": "",
                "description": "文档导入候选",
            }
            for item in _unique_keep_order(foreshadows, 40)
        ],
        "project_materials": {"bible": "", "world_rules": "", "timeline": "", "summary": ""},
    }


def _candidate_detail_text(kind, item):
    if not isinstance(item, dict):
        return ""
    if kind == "characters":
        lines = [
            f"姓名：{item.get('name', '')}",
            f"身份：{item.get('role', '')}",
            f"人物目标：{item.get('goal', '')}",
            f"隐藏秘密：{item.get('secret', '')}",
            f"语言风格：{item.get('voice', '')}",
            f"备注：{item.get('notes', '')}",
        ]
    elif kind == "lore":
        lines = [
            f"名称：{item.get('name', '')}",
            f"类型：{item.get('type', '')}",
            f"说明：{item.get('description', '')}",
        ]
    else:
        lines = [
            f"名称：{item.get('name', '')}",
            f"状态：{item.get('status', '')}",
            f"埋设章节：{item.get('setup_chapter', '')}",
            f"回收章节：{item.get('payoff_chapter', '')}",
            f"说明：{item.get('description', '')}",
        ]
    return "\n".join(str(x) for x in lines if str(x).strip())


def _candidate_field(item, *keys, default=""):
    item = item if isinstance(item, dict) else {}
    for key in keys:
        if key in item and str(item.get(key, "") or "").strip():
            return item.get(key)
    return default


def _candidate_group(data, *keys, default=None):
    data = data if isinstance(data, dict) else {}
    for key in keys:
        value = data.get(key)
        if value:
            return value
    return default


def _candidate_name(item):
    return str(_candidate_field(item, "name", "姓名", "名称", "名字", "伏笔名", "线索名", "标题") or "").strip()


_PROJECT_MATERIAL_ALIASES = {
    "bible": ("bible", "小说圣经", "圣经", "故事圣经", "作品圣经", "核心设定"),
    "world_rules": ("world_rules", "世界观", "规则", "世界规则", "设定规则", "世界观规则", "世界设定"),
    "timeline": ("timeline", "时间线", "时间顺序", "事件线", "剧情时间线"),
    "summary": ("summary", "摘要", "阶段摘要", "全局摘要", "故事摘要", "剧情摘要"),
}
_PROJECT_MATERIAL_KEYS = ("bible", "world_rules", "timeline", "summary")
_MANUAL_PROJECT_MATERIAL_KEYS = {"bible", "world_rules"}
_AUTO_PROJECT_MATERIAL_KEYS = tuple(
    key for key in _PROJECT_MATERIAL_KEYS if key not in _MANUAL_PROJECT_MATERIAL_KEYS
)


def _project_material_key(value):
    text = _compact_text(value).lower()
    if not text:
        return ""
    for key, aliases in _PROJECT_MATERIAL_ALIASES.items():
        for alias in aliases:
            alias_text = _compact_text(alias).lower()
            if text == alias_text or alias_text in text:
                return key
    return ""


def _normalize_ai_candidates(data):
    data = data if isinstance(data, dict) else {}
    out = {"characters": [], "lore": [], "foreshadows": [], "project_materials": {}}

    def merge_candidate_text(existing, incoming, append=False):
        old = str(existing or "").strip()
        new = str(incoming or "").strip()
        if append:
            cleaned_old = _dedupe_text_lines(old)
            if not new:
                return cleaned_old
            if not cleaned_old or cleaned_old in {"AI 分析候选", "文档导入候选"}:
                return _dedupe_text_lines(new)
            merged, _did_change = _merge_text_lines_without_duplicates(cleaned_old, new)
            return merged
        if not new:
            return old
        if not old or old in {"AI 分析候选", "文档导入候选"}:
            return new
        if new == old or new in old:
            return old
        if old in new:
            return new
        return old

    def merge_into_group(group_key, item):
        if not isinstance(item, dict):
            return
        name = str(item.get("name", "") or "").strip()
        if not name:
            return
        target = out[group_key]
        def merge_keys(value):
            keys = _record_alias_keys(value, name_key="name")
            if group_key == "characters":
                keys.add(_character_merge_key(value))
            return {key for key in keys if key}

        item_keys = merge_keys(item)
        existing = None
        for candidate in target:
            if merge_keys(candidate) & item_keys:
                existing = candidate
                break
        if existing is None:
            target.append(item)
            return
        if group_key == "foreshadows":
            status_probe = dict(existing)
            for field in ("status", "setup_chapter", "payoff_chapter", "description"):
                if str(item.get(field, "") or "").strip():
                    status_probe[field] = item.get(field, "")
            existing["status"] = _infer_foreshadow_status(status_probe, preserve_manual=False)
        if group_key == "characters" and name != str(existing.get("name", "") or "").strip():
            existing_name = str(existing.get("name", "") or "").strip()
            if existing_name:
                existing["notes"] = merge_candidate_text(existing.get("notes", ""), f"别称：{existing_name}", append=True)
            if _character_merge_key(item) == name:
                existing["name"] = name
            else:
                existing["notes"] = merge_candidate_text(existing.get("notes", ""), f"别称：{name}", append=True)
        elif group_key in {"lore", "foreshadows"} and name != str(existing.get("name", "") or "").strip():
            existing["description"] = merge_candidate_text(existing.get("description", ""), f"别称：{name}", append=True)
        for field, value in item.items():
            if field == "name" or (group_key == "foreshadows" and field == "status"):
                continue
            existing[field] = merge_candidate_text(
                existing.get(field, ""),
                value,
                append=field in {"notes", "description"},
            )

    def valid_items(key, *aliases):
        items = _candidate_group(data, key, *aliases, default=[])
        for item in items if isinstance(items, list) else []:
            name = _candidate_name(item)
            if isinstance(item, dict) and name:
                yield item, name

    for item, name in valid_items("characters", "人物", "人物卡", "角色", "角色列表", "character_candidates"):
        normalized_item = {
            "name": name,
            "role": str(_candidate_field(item, "role", "身份", "角色定位", "人物身份", "定位") or ""),
            "goal": str(_candidate_field(item, "goal", "目标", "人物目标", "动机", "诉求") or ""),
            "secret": str(_candidate_field(item, "secret", "秘密", "隐藏秘密", "隐情") or ""),
            "voice": str(_candidate_field(item, "voice", "语言风格", "说话风格", "口吻") or ""),
            "notes": str(_candidate_field(item, "notes", "备注", "说明", "关系变化", "人物关系", "description") or "AI 分析候选"),
        }
        char_item, lore_item = _prepare_character_candidate(normalized_item, default_notes="AI 分析候选")
        if char_item:
            merge_into_group("characters", char_item)
        elif lore_item:
            merge_into_group("lore", lore_item)
    for item, name in valid_items("lore", "设定", "设定库", "世界观设定", "资料设定", "settings"):
        merge_into_group("lore", {
            "name": name,
            "type": str(_candidate_field(item, "type", "类型", "类别", "设定类型") or "其他"),
            "description": str(_candidate_field(item, "description", "说明", "描述", "设定说明", "备注", "规则") or "AI 分析候选"),
        })
    for item, name in valid_items("foreshadows", "伏笔", "伏笔线索", "线索", "悬念", "clues"):
        normalized = {
            "name": name,
            "status": str(_candidate_field(item, "status", "状态", "伏笔状态") or "未埋"),
            "setup_chapter": str(_candidate_field(item, "setup_chapter", "埋设章节", "铺垫章节", "出现章节", "setup") or ""),
            "payoff_chapter": str(_candidate_field(item, "payoff_chapter", "回收章节", "兑现章节", "揭晓章节", "payoff") or ""),
            "description": str(_candidate_field(item, "description", "说明", "描述", "伏笔说明", "备注") or "AI 分析候选"),
        }
        normalized["status"] = _infer_foreshadow_status(normalized, preserve_manual=False)
        merge_into_group("foreshadows", normalized)
    materials = _candidate_group(data, "project_materials", "项目资料", "资料草案", "项目材料", "全局资料", default={})
    material_sources = []
    if isinstance(materials, dict):
        material_sources.append(materials)
    elif isinstance(materials, list):
        for item in materials:
            if not isinstance(item, dict):
                continue
            material_key = _project_material_key(
                _candidate_field(item, "key", "type", "类型", "字段", "名称", "标题", "category")
            )
            if not material_key:
                material_key = _project_material_key(
                    " ".join(str(k) for k in item.keys())
                )
            if not material_key:
                continue
            value = _candidate_field(item, "content", "内容", "text", "正文", "value", "说明", "description")
            if value:
                material_sources.append({material_key: value})
    material_sources.append(data)
    for source in material_sources:
        if isinstance(source, dict):
            for key in _PROJECT_MATERIAL_KEYS:
                out["project_materials"][key] = merge_candidate_text(
                    out["project_materials"].get(key, ""),
                    str(_candidate_field(source, *_PROJECT_MATERIAL_ALIASES[key]) or "").strip(),
                    append=True,
                )
    return out


def _apply_import_candidates(project, candidates, checked):
    project = project if isinstance(project, dict) else {}
    candidates = candidates if isinstance(candidates, dict) else {}
    checked = checked if isinstance(checked, dict) else {}
    result = {
        "added": {"characters": 0, "lore": 0, "foreshadows": 0},
        "merged": {"characters": 0, "lore": 0, "foreshadows": 0},
        "materials": {"bible": False, "world_rules": False, "timeline": False, "summary": False},
        "removed_candidates": 0,
    }

    chars = project.setdefault("characters", [])
    lore = project.setdefault("lore", [])
    foreshadows = project.setdefault("foreshadow_items", [])

    def useful(value):
        text = str(value or "").strip()
        return text and text not in {"AI 分析候选", "文档导入候选"}

    def merge_text(existing, incoming, append=False):
        old = str(existing or "").strip()
        new = str(incoming or "").strip()
        if append:
            cleaned_old = _dedupe_text_lines(old)
            if not useful(new):
                return cleaned_old, cleaned_old != old
            if not cleaned_old or cleaned_old in {"AI 分析候选", "文档导入候选"}:
                cleaned_new = _dedupe_text_lines(new)
                return cleaned_new, cleaned_new != old
            merged, did_change = _merge_text_lines_without_duplicates(cleaned_old, new)
            return merged, did_change or merged != old
        if not useful(new):
            return old, False
        if not old or old in {"AI 分析候选", "文档导入候选"}:
            return new, True
        if new == old or new in old:
            return old, False
        if append and old not in new:
            return f"{old}\n补充：{new}", True
        if len(new) > len(old) * 1.4 and old in new:
            return new, True
        return old, False

    def merge_item(target_item, incoming, fields, append_fields=()):
        changed = False
        for field in fields:
            merged, did_change = merge_text(
                target_item.get(field, ""),
                incoming.get(field, "") if isinstance(incoming, dict) else "",
                append=field in append_fields,
            )
            if did_change:
                target_item[field] = merged
                changed = True
        return changed

    def replace_linked_character(old_name, new_name):
        old_name = str(old_name or "").strip()
        new_name = str(new_name or "").strip()
        if not old_name or not new_name or old_name == new_name:
            return
        for chap in project.get("chapters", []) if isinstance(project.get("chapters", []), list) else []:
            if not isinstance(chap, dict) or not isinstance(chap.get("linked_characters"), list):
                continue
            linked = []
            seen = set()
            changed = False
            for value in chap.get("linked_characters", []):
                text = str(value or "").strip()
                if text == old_name:
                    text = new_name
                    changed = True
                if text and text not in seen:
                    seen.add(text)
                    linked.append(text)
            if changed:
                chap["linked_characters"] = linked

    def merge_key(key, item):
        if key == "characters":
            return _character_merge_key(item)
        return str(item.get("name", "") if isinstance(item, dict) else "").strip()

    def merge_keys(key, item):
        if not isinstance(item, dict):
            return set()
        keys = _record_alias_keys(item, name_key="name")
        if key == "characters":
            keys.add(_character_merge_key(item))
        return {value for value in keys if value}

    def apply_group(key, target, fields, append_fields=(), status_infer=None):
        existing = {}
        existing_aliases = {}
        duplicates = []
        merged_count = 0
        for x in target:
            if not isinstance(x, dict) or not str(x.get("name", "") or "").strip():
                continue
            if key == "characters":
                old_name, new_name = _canonicalize_character_record(x)
                replace_linked_character(old_name, new_name)
            item_key = merge_key(key, x)
            if not item_key:
                continue
            alias_keys = merge_keys(key, x) or {item_key}
            existing_item = None
            for alias_key in alias_keys:
                existing_item = existing_aliases.get(alias_key)
                if existing_item is not None:
                    break
            if existing_item is not None:
                alias_changed = False
                name = str(x.get("name", "") or "").strip()
                existing_name = str(existing_item.get("name", "") or "").strip()
                if name and name != existing_name:
                    if key == "characters":
                        existing_item["notes"], alias_changed = merge_text(
                            existing_item.get("notes", ""),
                            f"别称：{name}",
                            append=True,
                        )
                    elif key in {"lore", "foreshadows"}:
                        existing_item["description"], alias_changed = merge_text(
                            existing_item.get("description", ""),
                            f"别称：{name}",
                            append=True,
                        )
                if merge_item(existing_item, x, fields, append_fields):
                    merged_count += 1
                elif alias_changed:
                    merged_count += 1
                for alias_key in alias_keys:
                    existing_aliases[alias_key] = existing_item
                duplicates.append(x)
            else:
                existing[item_key] = x
                for alias_key in alias_keys:
                    existing_aliases[alias_key] = x
        for duplicate in duplicates:
            if duplicate in target:
                target.remove(duplicate)
        items = candidates.get(key, [])
        added_count = 0
        for i in checked.get(key, []):
            if i < 0 or i >= len(items):
                continue
            item = items[i]
            name = str(item.get("name", "") if isinstance(item, dict) else "").strip()
            if not name:
                continue
            item_key = merge_key(key, item)
            alias_keys = merge_keys(key, item) or {item_key}
            existing_item = None
            for alias_key in alias_keys:
                existing_item = existing_aliases.get(alias_key)
                if existing_item is not None:
                    break
            if existing_item is not None:
                if name != str(existing_item.get("name", "") or "").strip():
                    if key == "characters":
                        existing_item["notes"], alias_changed = merge_text(
                            existing_item.get("notes", ""),
                            f"别称：{name}",
                            append=True,
                        )
                    elif key in {"lore", "foreshadows"}:
                        existing_item["description"], alias_changed = merge_text(
                            existing_item.get("description", ""),
                            f"别称：{name}",
                            append=True,
                        )
                    else:
                        alias_changed = False
                else:
                    alias_changed = False
                changed = merge_item(existing_item, item, fields, append_fields) or alias_changed
                if status_infer is not None:
                    status_probe = deepcopy(existing_item)
                    for field in fields:
                        if useful(item.get(field, "") if isinstance(item, dict) else ""):
                            status_probe[field] = item.get(field, "")
                    inferred_status = status_infer(status_probe)
                    if inferred_status != str(existing_item.get("status", "") or "").strip():
                        existing_item["status"] = inferred_status
                        changed = True
                if changed:
                    merged_count += 1
                continue
            data = deepcopy(item)
            data.setdefault("id", uuid.uuid4().hex)
            if status_infer is not None:
                data["status"] = status_infer(data)
            target.append(data)
            existing[item_key] = data
            for alias_key in merge_keys(key, data) or {item_key}:
                existing_aliases[alias_key] = data
            added_count += 1
        return added_count, merged_count

    result["added"]["characters"], result["merged"]["characters"] = apply_group(
        "characters",
        chars,
        ("role", "goal", "secret", "voice", "notes"),
        append_fields=("notes",),
    )
    result["added"]["lore"], result["merged"]["lore"] = apply_group(
        "lore",
        lore,
        ("type", "description"),
        append_fields=("description",),
    )
    result["added"]["foreshadows"], result["merged"]["foreshadows"] = apply_group(
        "foreshadows",
        foreshadows,
        ("status", "setup_chapter", "payoff_chapter", "description"),
        append_fields=("description",),
        status_infer=lambda item: _infer_foreshadow_status(item, preserve_manual=False),
    )

    materials = candidates.get("project_materials", {})
    if isinstance(materials, dict):
        material_keys = _PROJECT_MATERIAL_KEYS
        selected = checked.get("project_materials", None)
        if isinstance(selected, list):
            selected_keys = set()
            for item in selected:
                if isinstance(item, int) and 0 <= item < len(material_keys):
                    selected_keys.add(material_keys[item])
                elif str(item) in material_keys:
                    selected_keys.add(str(item))
        else:
            selected_keys = set(_AUTO_PROJECT_MATERIAL_KEYS)
        for key in material_keys:
            if key not in selected_keys:
                continue
            merged, did_change = merge_text(project.get(key, ""), materials.get(key, ""), append=True)
            if did_change:
                project[key] = merged
                result["materials"][key] = True
                materials[key] = ""

    for key, indexes in checked.items():
        if key == "project_materials":
            continue
        items = candidates.get(key, [])
        for i in sorted(indexes, reverse=True):
            if 0 <= i < len(items):
                del items[i]
                result["removed_candidates"] += 1
    return result


def _candidate_dossier_detail_score(item, fields):
    item = item if isinstance(item, dict) else {}
    score = 0
    for field, weight in fields:
        value = str(item.get(field, "") or "").strip()
        if not value:
            continue
        compact = _compact_text(value)
        if compact in {"AI分析候选", "文档导入候选", "后续处理", "待后续处理", "待补充", "待定"}:
            continue
        score += weight + min(60, len(compact) // 4)
    return score


def _rank_candidate_dossier_characters(project, limit=40):
    project = project if isinstance(project, dict) else {}
    characters = project.get("characters", [])
    characters = characters if isinstance(characters, list) else []
    activity = _character_activity_stats(project)
    ranked = []
    for index, item in enumerate(characters):
        if not isinstance(item, dict):
            continue
        name = str(item.get("name", "") or "").strip()
        if not name:
            continue
        stat = activity.get(name, {})
        count = int(stat.get("chapters", 0) or 0)
        score = count * 1000
        score += _candidate_dossier_detail_score(
            item,
            (("role", 160), ("goal", 120), ("secret", 120), ("voice", 80), ("notes", 50)),
        )
        role_text = _compact_text(item.get("role", ""))
        notes_text = _compact_text(item.get("notes", ""))
        if any(key in role_text or key in notes_text for key in ("主角", "核心", "女主", "男主", "反派", "主线")):
            score += 500
        ranked.append((-score, -count, index, item))
    ranked.sort()
    return [item for _score, _count, _index, item in ranked[: max(0, int(limit or 0))]]


def _rank_candidate_dossier_lore(project, limit=40):
    project = project if isinstance(project, dict) else {}
    lore_items = project.get("lore", [])
    lore_items = lore_items if isinstance(lore_items, list) else []
    activity = _lore_activity_stats(project)
    ranked = []
    for index, item in enumerate(lore_items):
        if not isinstance(item, dict):
            continue
        name = str(item.get("name", "") or "").strip()
        if not name:
            continue
        stat = activity.get(name, {})
        count = int(stat.get("chapters", 0) or 0)
        score = count * 1000
        score += _candidate_dossier_detail_score(item, (("type", 80), ("description", 180)))
        type_text = _compact_text(item.get("type", ""))
        desc_text = _compact_text(item.get("description", ""))
        if any(key in type_text or key in desc_text for key in ("主线", "核心", "规则", "地点", "组织", "法则")):
            score += 250
        ranked.append((-score, -count, index, item))
    ranked.sort()
    return [item for _score, _count, _index, item in ranked[: max(0, int(limit or 0))]]


def _rank_candidate_dossier_foreshadows(project, limit=60):
    project = project if isinstance(project, dict) else {}
    items = project.get("foreshadow_items", [])
    items = items if isinstance(items, list) else []
    ranked = []
    status_weight = {"已埋": 5000, "未埋": 4500, "其他": 2400, "已回收": 900, "废弃": 0}
    for index, item in enumerate(items):
        if not isinstance(item, dict):
            continue
        name = str(item.get("name", "") or "").strip()
        if not name:
            continue
        status = _infer_foreshadow_status(item)
        score = status_weight.get(status, status_weight["其他"])
        if str(item.get("setup_chapter", "") or "").strip():
            score += 420
        if str(item.get("payoff_chapter", "") or "").strip():
            score += 360 if status in {"未埋", "已埋"} else 180
        score += _candidate_dossier_detail_score(item, (("description", 220),))
        desc_text = _compact_text(item.get("description", ""))
        if desc_text in {"AI分析候选", "文档导入候选", "后续处理", "待后续处理", "待回收"}:
            score -= 500
        ranked.append((-score, index if status in {"未埋", "已埋"} else -index, item))
    ranked.sort()
    return [item for _score, _tie_index, item in ranked[: max(0, int(limit or 0))]]


def _candidate_analysis_dossier_text(project):
    project = project if isinstance(project, dict) else {}
    context_parts = []
    meta = project.get("meta", {}) if isinstance(project.get("meta"), dict) else {}
    meta_lines = []
    for label, key in (("书名", "title"), ("类型", "genre"), ("风格", "style"), ("故事核心", "premise")):
        value = str(meta.get(key, "") or "").strip()
        if value:
            meta_lines.append(f"{label}：{value}")
    if meta_lines:
        context_parts.append("项目基础：\n" + "\n".join(meta_lines))
    for label, key, limit in (
        ("小说圣经", "bible", 1200),
        ("世界观/规则", "world_rules", 1000),
        ("时间线", "timeline", 1400),
        ("阶段摘要", "summary", 1400),
    ):
        value = _clip_context_text(_dedupe_text_lines(project.get(key, "")), limit, keep_tail=True)
        if value:
            context_parts.append(f"{label}：\n{value}")
    character_lines = []
    for item in _rank_candidate_dossier_characters(project, limit=40):
        if not isinstance(item, dict):
            continue
        name = str(item.get("name", "") or "").strip()
        if not name:
            continue
        detail = "｜".join(
            str(item.get(key, "") or "").strip()
            for key in ("role", "goal", "secret", "voice")
            if str(item.get(key, "") or "").strip()
        )
        character_lines.append(f"- {name}｜{detail}" if detail else f"- {name}")
    if character_lines:
        context_parts.append("已有主要人物（同名/别称请合并，不要重复新增）：\n" + "\n".join(character_lines))
    lore_lines = []
    for item in _rank_candidate_dossier_lore(project, limit=40):
        if not isinstance(item, dict):
            continue
        name = str(item.get("name", "") or "").strip()
        if not name:
            continue
        typ = str(item.get("type", "") or "其他").strip()
        desc = _clip_context_text(_dedupe_text_lines(item.get("description", "")), 120)
        lore_lines.append(f"- {name}｜{typ}｜{desc}" if desc else f"- {name}｜{typ}")
    if lore_lines:
        context_parts.append("已有设定库（同名/别称请合并，不要重复新增）：\n" + "\n".join(lore_lines))
    foreshadow_lines = []
    for item in _rank_candidate_dossier_foreshadows(project, limit=60):
        if not isinstance(item, dict):
            continue
        name = str(item.get("name", "") or "").strip()
        if not name:
            continue
        status = str(item.get("status", "") or "").strip()
        route = " -> ".join(
            str(item.get(key, "") or "").strip()
            for key in ("setup_chapter", "payoff_chapter")
            if str(item.get(key, "") or "").strip()
        )
        route = f"｜{route}" if route else ""
        desc = _clip_context_text(_dedupe_text_lines(item.get("description", "")), 120)
        tail = f"｜{desc}" if desc else ""
        foreshadow_lines.append(f"- {name}｜{status}{route}{tail}")
    if foreshadow_lines:
        context_parts.append("已有伏笔（继续补充状态/章节，不要重复新增）：\n" + "\n".join(foreshadow_lines))
    if not context_parts:
        return ""
    return "【已有项目档案】\n" + "\n\n".join(context_parts)


def _candidate_analysis_text(project, last_import_text="", chapter_ids=None, include_dossier=True):
    text = str(last_import_text or "").strip()
    if text and chapter_ids is None:
        return text
    project = project if isinstance(project, dict) else {}
    parts = []
    if include_dossier:
        dossier = _candidate_analysis_dossier_text(project)
        if dossier:
            parts.append(dossier)
    chapters = project.get("chapters", [])
    ids = {str(x) for x in chapter_ids} if chapter_ids is not None else None
    for chap in chapters if isinstance(chapters, list) else []:
        if not isinstance(chap, dict):
            continue
        if ids is not None and str(chap.get("id", "") or "") not in ids:
            continue
        parts.append(str(chap.get("title", "") or ""))
        outline = str(chap.get("outline", "") or "").strip()
        if outline:
            parts.append("章节提纲：\n" + outline)
        linked = chap.get("linked_characters", [])
        linked_names = _normalize_name_list(linked if isinstance(linked, list) else str(linked or ""))
        if linked_names:
            parts.append("关联人物：\n" + "、".join(linked_names))
        text_body = str(chap.get("text", "") or "").strip()
        if text_body:
            parts.append("正文：\n" + text_body)
        summary = str(chap.get("summary", "") or "").strip()
        if summary:
            parts.append("摘要：\n" + summary)
        key_facts = str(chap.get("key_facts", "") or "").strip()
        if key_facts:
            parts.append("关键事实：\n" + key_facts)
    return "\n".join(parts).strip()
