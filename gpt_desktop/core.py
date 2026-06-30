import sys
import os
import re
import json
import base64
import mimetypes
import uuid
import hashlib
import shutil
import subprocess
from datetime import datetime

from PySide6.QtCore import Qt, QUrl
from PySide6.QtGui import QPixmap, QPainter, QDesktopServices
from PySide6.QtWidgets import (
    QFileDialog,
)


# ============================================================
# 路径与常量
# ============================================================

APP_DIR = os.environ.get(
    "GPT_DESKTOP_APP_DIR",
    os.path.join(os.path.expanduser("~"), ".gpt_desktop_app"),
)
CONFIG_FILE = os.path.join(APP_DIR, "config.json")
HISTORY_DIR = os.path.join(APP_DIR, "history")
IMAGE_DIR = os.path.join(APP_DIR, "images")
VIDEO_DIR = os.path.join(APP_DIR, "videos")
IMAGE_UPLOAD_TMP_DIR = os.path.join(APP_DIR, "upload_tmp")
IMAGE_HISTORY_FILE = os.path.join(HISTORY_DIR, "image_history.json")
IMAGE_TASK_LOG_FILE = os.path.join(HISTORY_DIR, "image_task_log.json")
VIDEO_HISTORY_FILE = os.path.join(HISTORY_DIR, "video_history.json")
AGENT_HISTORY_FILE = os.path.join(HISTORY_DIR, "agent_history.json")
AGENT_SESSIONS_FILE = os.path.join(HISTORY_DIR, "agent_sessions.json")
INPUT_DRAFT_FILE = os.path.join(HISTORY_DIR, "input_drafts.json")
REFERENCE_DRAFT_DIR = os.path.join(HISTORY_DIR, "reference_image_drafts")
REFERENCE_SNAPSHOT_DIR = os.path.join(HISTORY_DIR, "reference_image_snapshots")

for d in (APP_DIR, HISTORY_DIR, IMAGE_DIR, VIDEO_DIR, IMAGE_UPLOAD_TMP_DIR, REFERENCE_DRAFT_DIR, REFERENCE_SNAPSHOT_DIR):
    os.makedirs(d, exist_ok=True)


# ============================================================
# 全局样式
# ============================================================

APP_STYLE = """
* { outline: 0; }

QMainWindow, QWidget {
    background-color: #16171c;
    color: #e8e8ea;
    font-family: "PingFang SC", "Helvetica Neue", Arial;
    font-size: 13px;
}

QLabel { color: #e8e8ea; }
QLabel#section_title { font-size: 20px; font-weight: 700; }
QLabel#sub_title    { font-size: 13px; font-weight: 600; color: #c8ccd6; }
QLabel#hint         { color: #8b8f99; font-size: 12px; }
QLabel#field_label  { color: #b6b9c2; }

QPushButton {
    background-color: #2a2c33;
    color: #e8e8ea;
    border: 1px solid #3a3c43;
    border-radius: 6px;
    padding: 6px 14px;
    min-height: 22px;
}
QPushButton:hover    { background-color: #353740; }
QPushButton:pressed  { background-color: #2a2c33; }
QPushButton:disabled { background-color: #1f2026; color: #6a6c72; border-color: #2a2c33; }

QPushButton#primary {
    background-color: #1f6feb;
    border-color: #1f6feb;
    color: white;
    font-weight: 600;
}
QPushButton#primary:hover    { background-color: #2c7cf0; }
QPushButton#primary:disabled { background-color: #2a3a55; color: #8a8d96; }

QPushButton#danger {
    background-color: transparent;
    border: 1px solid #b3434c;
    color: #d96a72;
}
QPushButton#danger:hover {
    background-color: #4a252b;
    border-color: #e06c75;
    color: #ff9aa2;
}
QPushButton#danger:pressed {
    background-color: #5a2a31;
    border-color: #ff7b86;
}

QPushButton#ghost { background-color: transparent; border: 1px solid #3a3c43; }
QPushButton#ghost:hover {
    background-color: #2b2f38;
    border-color: #5a6070;
    color: #ffffff;
}
QPushButton#ghost:pressed {
    background-color: #343947;
    border-color: #6b7284;
}
QPushButton#ghost:checked {
    background-color: #2b2f38;
    border-color: #6b7284;
    color: #ffffff;
}

QLineEdit, QTextEdit, QTextBrowser, QComboBox, QListWidget {
    background-color: #1a1b20;
    color: #e8e8ea;
    border: 1px solid #2a2c33;
    border-radius: 6px;
    padding: 6px 8px;
    selection-background-color: #1f6feb;
    selection-color: white;
}
QLineEdit:focus, QTextEdit:focus, QTextBrowser:focus, QComboBox:focus {
    border-color: #1f6feb;
}

QComboBox {
    padding-right: 28px;
}
QComboBox::drop-down {
    border: none;
    width: 26px;
    background: transparent;
}
QComboBox::down-arrow {
    image: none;
    width: 0px;
    height: 0px;
    border: none;
}
QComboBox QAbstractItemView {
    background-color: #1a1b20;
    color: #e8e8ea;
    border: 1px solid #2a2c33;
    selection-background-color: #1f6feb;
    selection-color: white;
    padding: 4px;
}
QComboBox QAbstractItemView::item { padding: 6px 10px; min-height: 22px; }

QListWidget::item { padding: 6px 4px; border-radius: 4px; }
QListWidget::item:selected { background-color: #1f6feb; color: white; }
QListWidget::item:hover    { background-color: #22242a; }

QTabWidget::pane { border: none; background: #16171c; top: -1px; }
QTabBar::tab {
    background: transparent;
    color: #8b8f99;
    padding: 10px 22px;
    border: none;
    font-size: 13px;
}
QTabBar::tab:selected { color: #ffffff; border-bottom: 2px solid #1f6feb; }
QTabBar::tab:hover { color: #e8e8ea; }

QFrame#card {
    background-color: #1a1b20;
    border: 1px solid #25272e;
    border-radius: 10px;
}

QScrollBar:vertical { background: transparent; width: 10px; margin: 2px; }
QScrollBar::handle:vertical { background: #3a3c43; border-radius: 5px; min-height: 30px; }
QScrollBar::handle:vertical:hover { background: #4a4c53; }
QScrollBar::add-line:vertical, QScrollBar::sub-line:vertical { height: 0; }
QScrollBar::add-page:vertical, QScrollBar::sub-page:vertical { background: transparent; }

QScrollBar:horizontal { background: transparent; height: 10px; margin: 2px; }
QScrollBar::handle:horizontal { background: #3a3c43; border-radius: 5px; min-width: 30px; }
QScrollBar::add-line:horizontal, QScrollBar::sub-line:horizontal { width: 0; }

QToolTip {
    background-color: #1a1b20;
    color: #e8e8ea;
    border: 1px solid #2a2c33;
    padding: 6px;
}
"""

# ============================================================
# 全局右键菜单反馈样式：
# 让输入框右键菜单有中文、hover 高亮、pressed 反馈
# ============================================================

CONTEXT_MENU_FEEDBACK_STYLE = """
QMenu {
    background-color: #111217;
    color: #f3f4f6;
    border: 1px solid #343741;
    border-radius: 8px;
    padding: 6px;
}

QMenu::item {
    padding: 7px 34px 7px 26px;
    min-width: 120px;
    border-radius: 5px;
    background-color: transparent;
}

QMenu::item:selected {
    background-color: #1f6feb;
    color: #ffffff;
}

QMenu::item:pressed {
    background-color: #1557b0;
    color: #ffffff;
}

QMenu::item:disabled {
    color: #6a6c72;
    background-color: transparent;
}

QMenu::separator {
    height: 1px;
    background-color: #30333b;
    margin: 6px 8px;
}
"""



# ============================================================
# 配置管理
# ============================================================


def safe_response_text(resp):
    """
    尽量避免接口错误信息乱码。
    优先按 utf-8 解码，失败后再用 requests 自己的 text。
    """
    try:
        return resp.content.decode("utf-8", errors="replace")
    except Exception:
        try:
            return resp.text
        except Exception:
            return ""


def extract_api_error(resp):
    """
    提取 API 错误信息，避免直接把整段 JSON 弹出来。
    """
    text = safe_response_text(resp)
    status_code = getattr(resp, "status_code", None)

    lower_text = text.lower()
    if status_code == 524 or ("error code 524" in lower_text and "cloudflare" in lower_text):
        return "中转服务器超时 524：Cloudflare 已连接到中转站，但中转站长时间没有返回结果。请稍后重试，或更换/联系该中转服务。"

    if "cloudflare" in lower_text and "a timeout occurred" in lower_text:
        return "中转服务器超时：Cloudflare 已连接到中转站，但中转站长时间没有返回结果。"

    try:
        data = json.loads(text)

        if isinstance(data, dict):
            err = data.get("error")
            if isinstance(err, dict):
                msg = err.get("message") or err.get("code") or err.get("type")
                if msg:
                    return str(msg)

            for key in ("message", "msg", "error_description", "detail"):
                if data.get(key):
                    return str(data.get(key))

        return text
    except Exception:
        return text


def fix_mojibake_text(text):
    """
    兜底修复常见 mojibake 乱码。
    例如：中文 UTF-8 被错误按 latin1/cp1252 解码后出现 å、æ、è 等。
    """
    if text is None:
        return ""

    text = str(text)

    suspicious_chars = ("å", "æ", "è", "é", "ä", "ç", "Ã", "Â", "¤", "¥", "½", "¾", "¼", "œ", "�")
    if not any(ch in text for ch in suspicious_chars):
        return text

    candidates = [text]

    try:
        candidates.append(text.encode("latin1").decode("utf-8"))
    except Exception:
        pass

    try:
        candidates.append(text.encode("latin1", errors="ignore").decode("utf-8", errors="ignore"))
    except Exception:
        pass

    try:
        candidates.append(text.encode("cp1252").decode("utf-8"))
    except Exception:
        pass

    try:
        candidates.append(text.encode("cp1252", errors="ignore").decode("utf-8", errors="ignore"))
    except Exception:
        pass

    def score(x):
        chinese = sum(1 for c in x if "\u4e00" <= c <= "\u9fff")
        bad = sum(x.count(ch) for ch in suspicious_chars)
        replacement = x.count("�")
        return chinese * 20 - bad * 5 - replacement * 10

    best = max(candidates, key=score)
    return best if score(best) > score(text) else text


def clean_error_text(err):
    """
    异常显示前的兜底清理。
    """
    text = fix_mojibake_text(str(err or "")).strip()
    return text or "请求失败，但没有返回具体错误信息。"


def log_debug(context, err=None):
    """低干扰调试日志：只输出到控制台，避免吞异常后完全无迹可查。"""
    try:
        if err is None:
            print(f"[debug] {context}")
        else:
            print(f"[debug] {context}: {err}")
    except Exception:
        pass


def make_clickable(widget, tooltip=None):
    """统一设置可点击控件的鼠标反馈。"""
    if widget is None:
        return widget
    widget.setCursor(Qt.PointingHandCursor)
    widget.setMouseTracking(True)
    widget.setAttribute(Qt.WA_Hover, True)
    if tooltip:
        widget.setToolTip(tooltip)
    try:
        widget.style().unpolish(widget)
        widget.style().polish(widget)
        widget.update()
    except Exception:
        pass
    return widget


def default_config():
    return {
        "providers": [],
        "model_cache": {},
        "image": {
            "provider_id": "",
            "model": "",
            "mode": "文生图",
            "size": "自动",
            "count": "1",
            "quality": "自动",
        },
        "video": {
            "provider_id": "",
            "model": "agnes-video-v2.0",
            "mode": "文生视频",
            "width": "1280",
            "height": "720",
            "num_frames": "81",
            "frame_rate": "24",
        },
        "agent": {"provider_id": "", "model": ""},
        "novel": {"provider_id": "", "model": ""},
    }


def normalize_config(cfg):
    if not isinstance(cfg, dict):
        return default_config()

    providers = cfg.get("providers")
    needs_migrate = providers is None
    if providers is None:
        providers = []
    elif not isinstance(providers, list):
        providers = []
    else:
        providers = [p for p in providers if isinstance(p, dict)]

    if needs_migrate:
        old_image = cfg.get("image", {}) or {}
        old_agent = cfg.get("agent", {}) or {}
        if not isinstance(old_image, dict):
            old_image = {}
        if not isinstance(old_agent, dict):
            old_agent = {}

        def maybe_add(name, base_url, api_key):
            if not (base_url or api_key):
                return ""
            for p in providers:
                if p.get("base_url") == base_url and p.get("api_key") == api_key:
                    return p["id"]
            pid = uuid.uuid4().hex[:8]
            providers.append({
                "id": pid, "name": name,
                "base_url": base_url or "", "api_key": api_key or "",
            })
            return pid

        img_pid = maybe_add("图片厂商", old_image.get("base_url", ""), old_image.get("api_key", ""))
        agt_pid = maybe_add("智能体厂商", old_agent.get("base_url", ""), old_agent.get("api_key", ""))
        cfg["image"] = {"provider_id": img_pid, "model": old_image.get("model", "")}
        cfg["agent"] = {"provider_id": agt_pid, "model": old_agent.get("model", "")}

    cfg["providers"] = providers
    if not isinstance(cfg.get("model_cache"), dict):
        cfg["model_cache"] = {}
    else:
        valid_ids = {p.get("id") for p in providers if p.get("id")}
        cfg["model_cache"] = {
            str(pid): [str(m) for m in models if str(m).strip()]
            for pid, models in cfg["model_cache"].items()
            if pid in valid_ids and isinstance(models, list)
        }
    if not isinstance(cfg.get("image"), dict):
        cfg["image"] = {}
    if not isinstance(cfg.get("video"), dict):
        cfg["video"] = {}
    if not isinstance(cfg.get("agent"), dict):
        cfg["agent"] = {}
    if not isinstance(cfg.get("novel"), dict):
        cfg["novel"] = {}
    cfg["image"].setdefault("provider_id", "")
    cfg["image"].setdefault("model", "")
    cfg["image"].setdefault("mode", "文生图")
    cfg["image"].setdefault("size", "自动")
    cfg["image"].setdefault("count", "1")
    cfg["image"].setdefault("quality", "自动")
    cfg["image"].setdefault("upload_optimization", "高质量")
    cfg["video"].setdefault("provider_id", "")
    cfg["video"].setdefault("model", "agnes-video-v2.0")
    cfg["video"].setdefault("mode", "文生视频")
    cfg["video"].setdefault("width", "1280")
    cfg["video"].setdefault("height", "720")
    cfg["video"].setdefault("num_frames", "81")
    cfg["video"].setdefault("frame_rate", "24")
    if not cfg["video"].get("provider_id") and cfg["image"].get("provider_id"):
        cfg["video"]["provider_id"] = cfg["image"].get("provider_id", "")
    cfg["agent"].setdefault("provider_id", "")
    cfg["agent"].setdefault("model", "")
    cfg["novel"].setdefault("provider_id", cfg["agent"].get("provider_id", ""))
    cfg["novel"].setdefault("model", cfg["agent"].get("model", ""))
    cfg["novel"].setdefault("candidate_analysis_concurrency", 3)
    cfg["novel"].setdefault("read_aloud_rate", "+0%")
    return cfg


def load_config():
    if not os.path.exists(CONFIG_FILE):
        return default_config()
    try:
        with open(CONFIG_FILE, "r", encoding="utf-8") as f:
            return normalize_config(json.load(f))
    except Exception:
        return default_config()


def _atomic_write_json(path, data):
    directory = os.path.dirname(path) or "."
    os.makedirs(directory, exist_ok=True)
    tmp_path = os.path.join(directory, f".{os.path.basename(path)}.{uuid.uuid4().hex}.tmp")
    try:
        with open(tmp_path, "w", encoding="utf-8") as f:
            json.dump(data, f, ensure_ascii=False, indent=2)
        os.replace(tmp_path, path)
    finally:
        try:
            if os.path.exists(tmp_path):
                os.remove(tmp_path)
        except Exception:
            pass


def save_config(cfg):
    try:
        _atomic_write_json(CONFIG_FILE, cfg)
    except Exception as e:
        print("保存配置失败：", e)


def load_json_file(path, default):
    if not os.path.exists(path):
        return default
    try:
        with open(path, "r", encoding="utf-8") as f:
            return json.load(f)
    except Exception:
        return default


def save_json_file(path, data):
    try:
        _atomic_write_json(path, data)
    except Exception as e:
        print("保存历史失败：", e)



def load_input_drafts():
    data = load_json_file(INPUT_DRAFT_FILE, {})
    return data if isinstance(data, dict) else {}


def save_input_drafts(data):
    if not isinstance(data, dict):
        data = {}
    save_json_file(INPUT_DRAFT_FILE, data)


def get_provider(config, provider_id):
    if not provider_id:
        return None
    for p in config.get("providers", []):
        if p.get("id") == provider_id:
            return p
    return None


def requests_proxies(proxy_url):
    proxy_url = str(proxy_url or "").strip()
    if not proxy_url:
        return None
    return {
        "http": proxy_url,
        "https": proxy_url,
    }


def api_url(base_url, path):
    base_url = (base_url or "").rstrip("/")
    if base_url.endswith("/v1") and path.startswith("/v1/"):
        return base_url + path[3:]
    return base_url + path


def now_str():
    return datetime.now().strftime("%Y-%m-%d %H:%M:%S")


def format_file_size(num_bytes):
    """
    把字节数格式化为 KB / MB。
    """
    try:
        num_bytes = int(num_bytes)
    except Exception:
        return ""

    if num_bytes < 1024:
        return f"{num_bytes}B"
    if num_bytes < 1024 * 1024:
        return f"{num_bytes / 1024:.0f}KB"
    return f"{num_bytes / 1024 / 1024:.1f}MB"


def hide_uploaded_file_content_for_display(text):
    """
    仅用于聊天窗口显示。

    目标显示效果：

    帮我修改这个程序……

    附件：
    main.py    93KB

    注意：
    这里只影响界面显示，不影响真正发送给智能体的内容。
    智能体仍然会收到完整文件内容。
    """
    if not isinstance(text, str):
        return text

    marker = "以下是用户上传的文件内容"
    if marker not in text:
        return text

    before, after = text.split(marker, 1)

    # 去掉类似：[已添加文件 1 个：main.py]
    # 支持它和用户输入在同一行的情况
    user_part = re.sub(
        r"\[已添加文件\s*\d+\s*个\s*[:：]\s*[^\]]+\]\s*",
        "",
        before
    ).strip()

    # 提取附件名和内容大小
    file_pattern = re.compile(r"【文件\s*[:：]\s*([^】]+)】")
    matches = list(file_pattern.finditer(after))

    attachments = []

    for i, m in enumerate(matches):
        filename = m.group(1).strip()

        content_start = m.end()
        content_end = matches[i + 1].start() if i + 1 < len(matches) else len(after)
        file_content = after[content_start:content_end]

        # 去掉前后的空白和代码块标记影响
        file_content = file_content.strip()

        size_text = format_file_size(len(file_content.encode("utf-8")))
        attachments.append((filename, size_text))

    parts = []

    if user_part:
        parts.append(user_part)

    if attachments:
        attach_lines = ["附件："]
        for filename, size_text in attachments:
            if size_text:
                attach_lines.append(f"{filename}    {size_text}")
            else:
                attach_lines.append(filename)
        parts.append("\n".join(attach_lines))
    else:
        parts.append("附件已发送给智能体。")

    return "\n\n".join(parts).strip()


def _dated_image_dir():
    """
    返回当天图片目录：
    ~/.gpt_desktop_app/images/YYYY-MM-DD
    """
    try:
        day = datetime.now().strftime("%Y-%m-%d")
        d = os.path.join(IMAGE_DIR, day)
        os.makedirs(d, exist_ok=True)
        return d
    except Exception:
        os.makedirs(IMAGE_DIR, exist_ok=True)
        return IMAGE_DIR


def _dated_video_dir():
    try:
        day = datetime.now().strftime("%Y-%m-%d")
        d = os.path.join(VIDEO_DIR, day)
        os.makedirs(d, exist_ok=True)
        return d
    except Exception:
        os.makedirs(VIDEO_DIR, exist_ok=True)
        return VIDEO_DIR


def image_suffix_from_content_type(content_type, default=".png"):
    text = str(content_type or "").split(";", 1)[0].strip().lower()
    mapping = {
        "image/png": ".png",
        "image/jpeg": ".jpg",
        "image/jpg": ".jpg",
        "image/webp": ".webp",
        "image/gif": ".gif",
        "image/bmp": ".bmp",
        "image/tiff": ".tiff",
    }
    return mapping.get(text, default)


def image_suffix_from_bytes(data_bytes, default=".png"):
    data = bytes(data_bytes or b"")[:16]
    if data.startswith(b"\x89PNG\r\n\x1a\n"):
        return ".png"
    if data.startswith(b"\xff\xd8\xff"):
        return ".jpg"
    if data.startswith(b"RIFF") and data[8:12] == b"WEBP":
        return ".webp"
    if data.startswith((b"GIF87a", b"GIF89a")):
        return ".gif"
    if data.startswith((b"II*\x00", b"MM\x00*")):
        return ".tiff"
    return default


def save_bytes_to_image(data_bytes, suffix=".png"):
    """
    新生成图片按日期分目录保存，避免 images 根目录无限堆积。
    """
    name = f"{uuid.uuid4().hex}{suffix}"
    path = os.path.join(_dated_image_dir(), name)
    with open(path, "wb") as f:
        f.write(data_bytes)
    return path


def save_video_response_to_file(resp, suffix=".mp4"):
    name = f"{uuid.uuid4().hex}{suffix}"
    path = os.path.join(_dated_video_dir(), name)
    with open(path, "wb") as f:
        for chunk in resp.iter_content(chunk_size=1024 * 1024):
            if chunk:
                f.write(chunk)
    if not os.path.exists(path) or os.path.getsize(path) <= 0:
        safe_remove_file(path)
        raise ValueError("视频文件保存后为空。")
    return path


def save_base64_to_image(b64_text, suffix=".png"):
    """
    把接口返回的 base64 图片直接分块写入文件，降低生成过程内存峰值。
    """
    text = str(b64_text or "").strip()
    if not text:
        raise ValueError("图片 base64 内容为空。")

    if "," in text and text.lower().startswith("data:image/"):
        text = text.split(",", 1)[1].strip()

    name = f"{uuid.uuid4().hex}{suffix}"
    path = os.path.join(_dated_image_dir(), name)

    with open(path, "wb") as f:
        chunk_size = 1024 * 1024
        for start in range(0, len(text), chunk_size):
            chunk = text[start:start + chunk_size]
            if len(chunk) % 4:
                next_start = start + chunk_size
                extra = text[next_start:next_start + (4 - len(chunk) % 4)]
                chunk += extra
            f.write(base64.b64decode(chunk))

    if not os.path.exists(path) or os.path.getsize(path) <= 0:
        safe_remove_file(path)
        raise ValueError("图片文件保存后为空。")

    return path


def image_file_to_base64(path):
    with open(path, "rb") as f:
        return base64.b64encode(f.read()).decode("utf-8")


def prepare_image_upload_file(path, mode="高质量"):
    """
    为图生图上传准备临时压缩副本。
    只返回上传用文件，不修改原图。返回：(path, mime, should_cleanup, message)
    """
    mode = str(mode or "高质量").strip()
    if mode == "关闭":
        return path, mimetypes.guess_type(path or "")[0] or "image/png", False, ""
    if mode == "标准":
        target_ratio = 0.50
        min_quality = 40
        start_quality = 98
    else:
        target_ratio = 0.70
        min_quality = 50
        start_quality = 98

    if not path or not os.path.exists(path):
        return path, mimetypes.guess_type(path or "")[0] or "image/png", False, ""

    original_mime = mimetypes.guess_type(path)[0] or "image/png"
    try:
        original_bytes = os.path.getsize(path)
    except Exception:
        original_bytes = 0

    try:
        from PySide6.QtGui import QImageReader, QImage

        reader = QImageReader(str(path))
        reader.setAutoTransform(True)

        img = reader.read()
        if img.isNull():
            return path, original_mime, False, ""

        if img.hasAlphaChannel():
            flattened = QImage(img.size(), QImage.Format_RGB888)
            flattened.fill(Qt.white)
            painter = QPainter(flattened)
            painter.drawImage(0, 0, img)
            painter.end()
            img = flattened
        else:
            img = img.convertToFormat(QImage.Format_RGB888)

        os.makedirs(IMAGE_UPLOAD_TMP_DIR, exist_ok=True)
        target_bytes = int(original_bytes * target_ratio) if original_bytes else 0
        best_path = ""
        best_bytes = 0
        best_quality = 0
        smallest_path = ""
        smallest_bytes = 0
        smallest_quality = 0

        for quality in range(int(start_quality), int(min_quality) - 1, -2):
            upload_path = os.path.join(IMAGE_UPLOAD_TMP_DIR, f"upload_{uuid.uuid4().hex}.jpg")
            if not img.save(upload_path, "JPEG", quality):
                safe_remove_file(upload_path)
                continue
            upload_bytes = os.path.getsize(upload_path) if os.path.exists(upload_path) else 0
            if upload_bytes <= 0:
                safe_remove_file(upload_path)
                continue

            if target_bytes and upload_bytes <= target_bytes:
                best_path = upload_path
                best_bytes = upload_bytes
                best_quality = quality
                break

            if not smallest_path or upload_bytes < smallest_bytes:
                safe_remove_file(smallest_path)
                smallest_path = upload_path
                smallest_bytes = upload_bytes
                smallest_quality = quality
            else:
                safe_remove_file(upload_path)

        if best_path and smallest_path and smallest_path != best_path:
            safe_remove_file(smallest_path)
        upload_path = best_path or smallest_path
        upload_bytes = best_bytes or smallest_bytes
        jpeg_quality = best_quality or smallest_quality
        if not upload_path:
            return path, original_mime, False, ""

        try:
            upload_bytes = os.path.getsize(upload_path)
        except Exception:
            upload_bytes = 0

        if upload_bytes <= 0:
            safe_remove_file(upload_path)
            return path, original_mime, False, ""

        if original_bytes and upload_bytes >= original_bytes:
            safe_remove_file(upload_path)
            return path, original_mime, False, ""

        ratio_text = f"{int((upload_bytes / original_bytes) * 100)}%" if original_bytes else ""
        msg = f"{os.path.basename(path)}：{format_file_size(original_bytes)} -> {format_file_size(upload_bytes)}（约{ratio_text}，质量 {jpeg_quality}）"
        return upload_path, "image/jpeg", True, msg
    except Exception:
        return path, original_mime, False, ""


def read_uploaded_files_text(file_paths):
    MAX_UPLOAD_FILE_SIZE = 10 * 1024 * 1024

    if not file_paths:
        return ""

    def _format_size(n):
        if n >= 1024 * 1024:
            return f"{n / (1024 * 1024):.2f} MB"
        if n >= 1024:
            return f"{n / 1024:.2f} KB"
        return f"{n} bytes"

    def _decode_text_bytes(raw):
        for enc in ("utf-8", "utf-8-sig", "gb18030", "gbk", "big5", "latin1"):
            try:
                return raw.decode(enc)
            except Exception:
                pass
        return raw.decode("utf-8", errors="replace")

    def _read_plain_text(file_path):
        try:
            with open(file_path, "rb") as f:
                raw = f.read()
            return _decode_text_bytes(raw), ""
        except Exception as e:
            return "", f"文本文件读取失败：{e}"

    def _read_docx_text(file_path):
        try:
            from docx import Document
        except Exception:
            return "", "未安装 DOCX 解析库，请执行：pip install python-docx"

        try:
            doc = Document(file_path)
            chunks = []

            for p in doc.paragraphs:
                t = (p.text or "").strip()
                if t:
                    chunks.append(t)

            for table in doc.tables:
                for row in table.rows:
                    cells = []
                    for cell in row.cells:
                        t = (cell.text or "").strip()
                        if t:
                            cells.append(t)
                    if cells:
                        chunks.append(" | ".join(cells))

            content = "\n\n".join(chunks).strip()
            if not content:
                return "", "DOCX 文件中没有提取到可读文本。"
            return content, ""
        except Exception as e:
            return "", f"DOCX 解析失败：{e}"

    def _read_pdf_text(file_path):
        try:
            from pypdf import PdfReader
        except Exception:
            return "", "未安装 PDF 解析库，请执行：pip install pypdf"

        try:
            reader = PdfReader(file_path)
            chunks = []
            for page in reader.pages:
                try:
                    t = page.extract_text() or ""
                except Exception:
                    t = ""
                t = t.strip()
                if t:
                    chunks.append(t)

            content = "\n\n".join(chunks).strip()
            if not content:
                return "", "PDF 文件中没有提取到可读文本。"
            return content, ""
        except Exception as e:
            return "", f"PDF 解析失败：{e}"

    parts = []
    max_chars_per_file = 120000

    for file_path in file_paths:
        try:
            file_path = str(file_path)

            if not file_path or not os.path.exists(file_path):
                parts.append(f"\n\n[附件文件缺失]\n路径：{file_path}\n")
                continue

            name = os.path.basename(file_path)
            size = os.path.getsize(file_path)
            ext = os.path.splitext(file_path)[1].lower()

            if size > MAX_UPLOAD_FILE_SIZE:
                parts.append(
                    f"\n\n[附件文件过大]\n"
                    f"文件名：{name}\n"
                    f"文件大小：{_format_size(size)}\n"
                    f"限制：10 MB\n"
                    f"提示：请压缩、拆分或另存为更小文件后再上传。\n"
                )
                continue

            content = ""
            note = ""

            if ext == ".docx":
                content, note = _read_docx_text(file_path)
            elif ext == ".pdf":
                content, note = _read_pdf_text(file_path)
            elif ext == ".doc":
                note = "暂不支持旧版 .doc 文件，请先用 Word/WPS 另存为 .docx 后再上传。"
            else:
                content, note = _read_plain_text(file_path)

            if content and len(content) > max_chars_per_file:
                content = content[:max_chars_per_file] + "\n...[内容过长，已截断]"

            block = (
                f"\n\n[附件文件]\n"
                f"文件名：{name}\n"
                f"文件大小：{_format_size(size)}\n"
            )

            if note:
                block += f"解析提示：{note}\n"

            if content:
                block += "文件内容如下：\n```\n" + content + "\n```\n"
            elif not note:
                block += "文件内容为空或无法读取。\n"

            parts.append(block)
        except Exception as e:
            parts.append(
                f"\n\n[附件读取失败]\n"
                f"路径：{file_path}\n"
                f"错误：{e}\n"
            )

    return "".join(parts)



# ============================================================
# 唯一缩略图缓存方案
#
# 说明：
# 1. 缩略图缓存按：原图绝对路径 + 文件大小 + 修改时间 + 目标尺寸 生成唯一 key；
# 2. UI 里 load_thumbnail_pixmap() 只读取小图缓存，避免加载原图卡顿；
# ============================================================

THUMBNAIL_CACHE_DIR = os.path.join(APP_DIR, "thumbnail_cache")


def thumbnail_cache_dir():
    try:
        os.makedirs(THUMBNAIL_CACHE_DIR, exist_ok=True)
        return THUMBNAIL_CACHE_DIR
    except Exception:
        os.makedirs(APP_DIR, exist_ok=True)
        return APP_DIR


def thumbnail_cache_path(path, max_w=220, max_h=220):
    """
    根据原图状态和目标尺寸生成唯一缩略图缓存路径。
    """
    try:
        p = os.path.abspath(str(path))
        st = os.stat(p)
        raw = json.dumps(
            {
                "path": p,
                "size": int(st.st_size),
                "mtime_ns": int(getattr(st, "st_mtime_ns", int(st.st_mtime * 1000000000))),
                "w": int(max_w),
                "h": int(max_h),
            },
            ensure_ascii=False,
            sort_keys=True,
        )
    except Exception:
        raw = json.dumps(
            {
                "path": str(path),
                "w": int(max_w),
                "h": int(max_h),
            },
            ensure_ascii=False,
            sort_keys=True,
        )

    name = hashlib.sha256(raw.encode("utf-8", errors="ignore")).hexdigest() + ".png"
    return os.path.join(thumbnail_cache_dir(), name)


def ensure_thumbnail_cache(path, max_w=220, max_h=220):
    """
    确保某个尺寸的缩略图缓存存在。

    可在后台线程调用。
    注意：这里使用 QImageReader / QImage，不创建 QPixmap。
    """
    try:
        if not path or not os.path.exists(path):
            return ""

        max_w = int(max_w)
        max_h = int(max_h)
        cache_path = thumbnail_cache_path(path, max_w, max_h)

        if os.path.exists(cache_path) and os.path.getsize(cache_path) > 0:
            return cache_path

        from PySide6.QtGui import QImageReader
        from PySide6.QtCore import QSize

        reader = QImageReader(str(path))
        reader.setAutoTransform(True)

        # 尽量让 Qt 按目标尺寸解码，减少大图解码压力。
        try:
            original_size = reader.size()
            if original_size.isValid() and original_size.width() > 0 and original_size.height() > 0:
                scaled_size = original_size.scaled(
                    QSize(max_w, max_h),
                    Qt.KeepAspectRatio,
                )
                if scaled_size.isValid():
                    reader.setScaledSize(scaled_size)
        except Exception:
            pass

        img = reader.read()
        if img.isNull():
            return ""

        # 兜底限制尺寸。
        if img.width() > max_w or img.height() > max_h:
            img = img.scaled(
                max_w,
                max_h,
                Qt.KeepAspectRatio,
                Qt.SmoothTransformation,
            )

        tmp_path = cache_path + "." + uuid.uuid4().hex + ".tmp"
        if not img.save(tmp_path, "PNG"):
            try:
                if os.path.exists(tmp_path):
                    os.remove(tmp_path)
            except Exception:
                pass
            return ""

        try:
            os.replace(tmp_path, cache_path)
        except Exception:
            shutil.move(tmp_path, cache_path)

        return cache_path

    except Exception as e:
        try:
            print("生成缩略图缓存失败：", e)
        except Exception:
            pass
        return ""


def load_thumbnail_pixmap(path, max_w=220, max_h=220, generate_missing=True):
    """
    加载缩略图 QPixmap。

    - 优先读取唯一缩略图缓存；
    - generate_missing=True 时会生成缺失缓存；
    - 图库启动渲染应使用 generate_missing=False，避免 UI 线程解码原图；
    - 避免在主线程直接加载大图再缩放。
    """
    try:
        max_w = int(max_w)
        max_h = int(max_h)

        cache_path = thumbnail_cache_path(path, max_w, max_h)
        if (not os.path.exists(cache_path) or os.path.getsize(cache_path) <= 0) and generate_missing:
            cache_path = ensure_thumbnail_cache(path, max_w, max_h)

        if cache_path and os.path.exists(cache_path):
            pix = QPixmap(str(cache_path))
            if not pix.isNull():
                return pix

    except Exception:
        pass

    return QPixmap()


def ensure_video_thumbnail_cache(path, max_w=420, max_h=280):
    """
    为本地视频生成封面缩略图缓存。

    macOS 上优先使用系统 QuickLook，不依赖 ffmpeg。失败时返回空字符串，
    UI 会显示非黑色占位，避免历史区域一片黑。
    """
    try:
        if not path or not os.path.exists(path):
            return ""

        max_w = int(max_w)
        max_h = int(max_h)
        cache_path = thumbnail_cache_path(path, max_w, max_h)
        if os.path.exists(cache_path) and os.path.getsize(cache_path) > 0:
            return cache_path

        tmp_dir = os.path.join(thumbnail_cache_dir(), ".video_" + uuid.uuid4().hex)
        os.makedirs(tmp_dir, exist_ok=True)
        try:
            if sys.platform == "darwin":
                subprocess.run(
                    ["qlmanage", "-t", "-s", str(max(max_w, max_h)), "-o", tmp_dir, path],
                    stdout=subprocess.DEVNULL,
                    stderr=subprocess.DEVNULL,
                    timeout=12,
                    check=False,
                )

            candidates = []
            for name in os.listdir(tmp_dir):
                fp = os.path.join(tmp_dir, name)
                if os.path.isfile(fp) and os.path.getsize(fp) > 0:
                    candidates.append(fp)

            if not candidates:
                return ""

            src = max(candidates, key=lambda p: os.path.getsize(p))
            pix = QPixmap(src)
            if pix.isNull():
                return ""
            pix = pix.scaled(max_w, max_h, Qt.KeepAspectRatio, Qt.SmoothTransformation)

            tmp_path = cache_path + "." + uuid.uuid4().hex + ".tmp"
            if not pix.save(tmp_path, "PNG"):
                safe_remove_file(tmp_path)
                return ""
            try:
                os.replace(tmp_path, cache_path)
            except Exception:
                shutil.move(tmp_path, cache_path)
            return cache_path
        finally:
            try:
                shutil.rmtree(tmp_dir)
            except Exception:
                pass
    except Exception:
        return ""


def load_video_thumbnail_pixmap(path, max_w=420, max_h=280, generate_missing=True):
    try:
        cache_path = thumbnail_cache_path(path, max_w, max_h)
        if (not os.path.exists(cache_path) or os.path.getsize(cache_path) <= 0) and generate_missing:
            cache_path = ensure_video_thumbnail_cache(path, max_w, max_h)
        if cache_path and os.path.exists(cache_path):
            pix = QPixmap(str(cache_path))
            if not pix.isNull():
                return pix
    except Exception:
        pass
    return QPixmap()


def get_path_size(path):
    """
    递归统计文件/目录占用大小。
    """
    total = 0
    try:
        if not path or not os.path.exists(path):
            return 0

        if os.path.isfile(path):
            return os.path.getsize(path)

        for root, dirs, files in os.walk(path):
            for name in files:
                fp = os.path.join(root, name)
                try:
                    total += os.path.getsize(fp)
                except Exception:
                    pass
    except Exception:
        pass
    return total


def format_cache_size(num_bytes):
    """
    缓存大小显示。
    """
    try:
        num_bytes = int(num_bytes)
    except Exception:
        return "0 B"

    if num_bytes < 1024:
        return f"{num_bytes} B"
    if num_bytes < 1024 ** 2:
        return f"{num_bytes / 1024:.1f} KB"
    if num_bytes < 1024 ** 3:
        return f"{num_bytes / 1024 ** 2:.1f} MB"
    return f"{num_bytes / 1024 ** 3:.2f} GB"


def safe_clear_dir(path):
    """
    清空目录内容，但保留目录本身。
    """
    failures = []
    try:
        os.makedirs(path, exist_ok=True)
        for name in os.listdir(path):
            fp = os.path.join(path, name)
            try:
                if os.path.isdir(fp):
                    shutil.rmtree(fp)
                else:
                    os.remove(fp)
            except Exception as e:
                failures.append(f"{fp}: {e}")
    except Exception as e:
        failures.append(f"{path}: {e}")
    return failures


def safe_remove_file(path):
    """
    安全删除单个文件。
    """
    try:
        if path and os.path.exists(path) and os.path.isfile(path):
            os.remove(path)
        return True, ""
    except Exception as e:
        return False, str(e)
    return True, ""


def open_local_file(path):
    try:
        if not path or not os.path.exists(path):
            return False
        if sys.platform == "darwin":
            subprocess.Popen(["open", path])
            return True
        return QDesktopServices.openUrl(QUrl.fromLocalFile(path))
    except Exception:
        return False


def _escape_windows_batch_value(value):
    return str(value or "").replace("%", "%%")


def _build_windows_open_after_exit_script(target_path, current_pid, process_name, max_wait_seconds=120):
    target_path = _escape_windows_batch_value(os.path.abspath(target_path))
    process_name = _escape_windows_batch_value(process_name)
    return "\r\n".join([
        "@echo off",
        "setlocal EnableExtensions EnableDelayedExpansion",
        f"set \"TARGET={target_path}\"",
        f"set \"APP_PID={int(current_pid)}\"",
        f"set \"APP_PROCESS={process_name}\"",
        f"set \"MAX_WAIT={int(max_wait_seconds)}\"",
        "set /a WAITED=0",
        ":wait_pid",
        "tasklist /FI \"PID eq %APP_PID%\" 2>NUL | find \"%APP_PID%\" >NUL",
        "if not errorlevel 1 (",
        "  timeout /T 1 /NOBREAK >NUL",
        "  set /a WAITED+=1",
        "  if !WAITED! LSS !MAX_WAIT! goto wait_pid",
        ")",
        "set /a WAITED=0",
        ":wait_process",
        "if not \"%APP_PROCESS%\"==\"\" (",
        "  tasklist /FI \"IMAGENAME eq %APP_PROCESS%\" 2>NUL | find /I \"%APP_PROCESS%\" >NUL",
        "  if not errorlevel 1 (",
        "    timeout /T 1 /NOBREAK >NUL",
        "    set /a WAITED+=1",
        "    if !WAITED! LSS 10 goto wait_process",
        "  )",
        ")",
        "start \"\" \"%TARGET%\"",
        "del \"%~f0\" >NUL 2>NUL",
        "",
    ])


def open_local_file_after_app_exit(path):
    try:
        if not path or not os.path.exists(path):
            return False

        if sys.platform.startswith("win"):
            updates_dir = os.path.join(APP_DIR, "updates")
            os.makedirs(updates_dir, exist_ok=True)
            script_path = os.path.join(updates_dir, f"open_update_after_exit_{os.getpid()}.cmd")
            process_name = os.path.basename(sys.executable or "")
            script = _build_windows_open_after_exit_script(path, os.getpid(), process_name)
            with open(script_path, "w", encoding="utf-8", newline="") as f:
                f.write(script)

            flags = 0
            for name in ("CREATE_NO_WINDOW", "CREATE_NEW_PROCESS_GROUP"):
                flags |= int(getattr(subprocess, name, 0))
            subprocess.Popen(
                ["cmd.exe", "/c", script_path],
                stdin=subprocess.DEVNULL,
                stdout=subprocess.DEVNULL,
                stderr=subprocess.DEVNULL,
                close_fds=True,
                creationflags=flags,
            )
            return True

        if sys.platform == "darwin":
            subprocess.Popen(
                ["/bin/sh", "-c", "while kill -0 \"$1\" 2>/dev/null; do sleep 0.5; done; open \"$2\"", "open-after-exit", str(os.getpid()), path],
                stdin=subprocess.DEVNULL,
                stdout=subprocess.DEVNULL,
                stderr=subprocess.DEVNULL,
                close_fds=True,
            )
            return True

        subprocess.Popen(
            ["/bin/sh", "-c", "while kill -0 \"$1\" 2>/dev/null; do sleep 0.5; done; xdg-open \"$2\"", "open-after-exit", str(os.getpid()), path],
            stdin=subprocess.DEVNULL,
            stdout=subprocess.DEVNULL,
            stderr=subprocess.DEVNULL,
            close_fds=True,
        )
        return True
    except Exception:
        return False


# ============================================================
# 中文文件选择器
# ============================================================

def get_open_file_names_cn(parent, title, file_filter, start_dir=""):
    """
    使用 macOS 系统原生 Finder 文件选择器。

    目标效果：
    - macOS 下显示类似 Finder 的上传/打开窗口；
    - 左侧有 最近使用、应用程序、桌面、文稿、下载、iCloud 云盘等；
    - 默认打开“下载”目录；
    - 不使用 Qt 自绘文件选择器。

    重点：
    """
    if start_dir:
        directory = start_dir
    else:
        download_dir = os.path.join(os.path.expanduser("~"), "Downloads")
        directory = download_dir if os.path.exists(download_dir) else os.path.expanduser("~")

    options = QFileDialog.Options()

    files, _ = QFileDialog.getOpenFileNames(
        parent,
        title or "打开",
        directory,
        file_filter,
        "",
        options
    )
    return files or []


def get_save_file_name_cn(parent, title, default_name, file_filter, start_dir=""):
    """
    使用系统原生保存文件选择器。
    """
    base_dir = start_dir or os.path.expanduser("~")
    default_path = os.path.join(base_dir, default_name) if default_name else base_dir

    file_path, _ = QFileDialog.getSaveFileName(
        parent,
        title or "保存文件",
        default_path,
        file_filter,
        "",
        QFileDialog.Options()
    )
    return file_path or ""


# ============================================================
# Markdown 渲染（极简）
# ============================================================

CHAT_CODE_COPY_STORE = {}

def looks_like_code_block(language, code):
    language = (language or "").strip().lower()
    code = code or ""
    if language:
        return language not in ("text", "txt", "plain", "plaintext", "中文", "说明")

    stripped = code.strip()
    if not stripped:
        return False

    lines = [line.strip() for line in stripped.splitlines() if line.strip()]
    if not lines:
        return False

    code_markers = (
        "def ", "class ", "import ", "from ", "return ", "function ", "const ", "let ", "var ",
        "if ", "for ", "while ", "try:", "except ", "public ", "private ", "#include",
        "SELECT ", "INSERT ", "UPDATE ", "DELETE ", "CREATE ", "curl ", "python ", "pip ",
        "npm ", "yarn ", "git ", "docker ", "brew ", "cd ", "ls ", "chmod ", "sudo ",
    )
    symbol_markers = ("{", "}", "</", "/>", "=>", "==", "!=", "::", "&&", "||", "$(", "#!")

    joined = "\n".join(lines[:12])
    upper_joined = joined.upper()
    marker_hits = sum(1 for marker in code_markers if marker in joined or marker in upper_joined)
    symbol_hits = sum(1 for marker in symbol_markers if marker in joined)

    if marker_hits or symbol_hits >= 2:
        return True

    command_like = sum(1 for line in lines[:8] if re.match(r"^[$>#]?\s*[\w./-]+(\s+[-\w./:=]+)+$", line))
    return command_like >= 2

def normalize_chat_text_newlines(text):
    if not text:
        return ""

    text = text.replace("\r\n", "\n").replace("\r", "\n")
    lines = [line.rstrip() for line in text.split("\n")]
    normalized = [line for line in lines if line.strip()]
    return "\n".join(normalized).strip()

def md_to_html(text):
    if not text:
        return ""

    code_blocks = []
    link_blocks = []

    def stash_code(match):
        language = match.group(1)
        code = match.group(2)
        safe = code.replace("&", "&amp;").replace("<", "&lt;").replace(">", "&gt;")
        if not looks_like_code_block(language, code):
            html_text = normalize_chat_text_newlines(safe).replace("\n", "<br>")
            code_blocks.append(html_text)
            return f"\x00CODE{len(code_blocks) - 1}\x00"

        code_id = uuid.uuid4().hex
        CHAT_CODE_COPY_STORE[code_id] = code.rstrip("\n")
        html_text = (
            '<table width="100%" cellspacing="0" cellpadding="0" border="0" '
            'bgcolor="#1f2229" style="border:1px solid #343843; margin:8px 0;">'
            '<tr><td style="padding:7px 12px; border-bottom:1px solid #343843;" align="right">'
            f'<a href="app://copy-code?id={code_id}" style="color:#8ab4f8;text-decoration:none;font-size:12px;">复制代码</a>'
            '</td></tr>'
            '<tr><td style="padding:10px 12px;">'
            '<pre style="margin:0; background:transparent; color:#e6e8ee;'
            'font-family:Menlo,Monaco,monospace;'
            'white-space:pre-wrap; font-size:12px; line-height:1.55;">'
            f'{safe}'
            '</pre>'
            '</td></tr></table>'
        )
        code_blocks.append(html_text)
        return f"\x00CODE{len(code_blocks) - 1}\x00"

    def make_link(label, url):
        label = (label or "").strip()
        url = (url or "").strip()
        if not label or not url:
            return label

        href = url.replace("&amp;", "&")
        if not re.match(r"^[a-zA-Z][a-zA-Z0-9+.-]*:", href):
            href = "https://" + href

        safe_href = href.replace("&", "&amp;").replace('"', "&quot;").replace("<", "%3C").replace(">", "%3E")
        safe_label = label.replace('"', "&quot;")
        return f'<a href="{safe_href}" style="color:#58a6ff;text-decoration:none;">{safe_label}</a>'

    def stash_inline_code(match):
        code = match.group(1)
        html_text = (
            f'<code style="background-color:#242733;color:#f0f1f3;padding:1px 5px;border-radius:3px;'
            f'font-family:Menlo,Monaco,monospace;">{code}</code>'
        )
        code_blocks.append(html_text)
        return f"\x00CODE{len(code_blocks) - 1}\x00"

    def apply_links(value):
        value = re.sub(
            r"\[([^\]\n]+)\]\((https?://[^\s\)]+)\)",
            lambda m: stash_link(make_link(m.group(1), m.group(2))),
            value,
        )

        def replace_bare_url(match):
            url = match.group(0)
            tail = ""
            while url and url[-1] in ".,;:!?":
                tail = url[-1] + tail
                url = url[:-1]
            while url.endswith(")") and url.count("(") < url.count(")"):
                tail = ")" + tail
                url = url[:-1]
            return make_link(url, url) + tail

        return re.sub(r"(?<![\"'=])(https?://[^\s<]+)", replace_bare_url, value)

    def stash_link(html_text):
        link_blocks.append(html_text)
        return f"\x00LINK{len(link_blocks) - 1}\x00"

    text = re.sub(r"```(\w*)\n?(.*?)```", stash_code, text, flags=re.DOTALL)
    text = normalize_chat_text_newlines(text)
    text = text.replace("&", "&amp;").replace("<", "&lt;").replace(">", "&gt;")
    text = re.sub(r"`([^`\n]+)`", stash_inline_code, text)
    text = apply_links(text)

    lines = text.split("\n")
    out = []
    in_list = False

    def close_list():
        nonlocal in_list
        if in_list:
            out.append("</ul>")
            in_list = False

    for line in lines:
        s = line.lstrip()
        if s.startswith("### "):
            close_list()
            out.append(f'<div style="font-weight:700;font-size:14px;margin:8px 0 2px;">{s[4:]}</div>')
        elif s.startswith("## "):
            close_list()
            out.append(f'<div style="font-weight:700;font-size:16px;margin:8px 0 2px;">{s[3:]}</div>')
        elif s.startswith("# "):
            close_list()
            out.append(f'<div style="font-weight:700;font-size:18px;margin:8px 0 2px;">{s[2:]}</div>')
        elif s.startswith("- ") or s.startswith("* "):
            if not in_list:
                out.append('<ul style="margin:4px 0 4px 18px;">')
                in_list = True
            out.append(f"<li>{s[2:]}</li>")
        elif s == "---":
            close_list()
            out.append('<hr style="border:0;border-top:1px solid #3a3c43;margin:8px 0;">')
        else:
            close_list()
            out.append(line)
    close_list()
    text = "\n".join(out)

    text = re.sub(r"\*\*(.+?)\*\*", r"<b>\1</b>", text)
    text = re.sub(r"(?<![\*\w])\*([^\*\n]+)\*(?!\*)", r"<i>\1</i>", text)
    text = text.replace("\n", "<br>")
    text = re.sub(r"(<ul[^>]*>)<br>", r"\1", text)
    text = text.replace("<br><li>", "<li>")
    text = text.replace("</li><br><li>", "</li><li>")
    text = text.replace("</li><br></ul>", "</li></ul>")
    text = text.replace("<br></ul>", "</ul>")

    def restore_code(match):
        return code_blocks[int(match.group(1))]

    def restore_link(match):
        return link_blocks[int(match.group(1))]

    text = re.sub(r"\x00CODE(\d+)\x00", restore_code, text)
    return re.sub(r"\x00LINK(\d+)\x00", restore_link, text)


# ============================================================
# 后台线程
# ============================================================
